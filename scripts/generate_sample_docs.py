"""Create sample PDF, Excel, Word files from CSV for demo uploads."""

from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path

import pandas as pd
from fpdf import FPDF


def make_excel(csv_path: Path, out_path: Path) -> None:
    df = pd.read_csv(csv_path)
    df.to_excel(out_path, index=False, sheet_name="Data")


def make_pdf_from_csv(csv_path: Path, out_path: Path, title: str) -> None:
    """PDF as structured text lines — reliable extract, strict columns on read."""
    df = pd.read_csv(csv_path)
    pdf = FPDF(format="A4", orientation="L")
    pdf.set_margins(8, 8, 8)
    pdf.set_auto_page_break(auto=True, margin=8)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 6)
    width = pdf.epw
    for _, row in df.iterrows():
        parts = [f"{col}={row[col]}" for col in df.columns]
        line = " | ".join(parts)
        safe = line.encode("latin-1", errors="replace").decode("latin-1")
        pdf.multi_cell(width, 3, safe)

    pdf.output(out_path)


def make_docx_from_csv(csv_path: Path, out_path: Path, title: str) -> None:
    from docx import Document

    df = pd.read_csv(csv_path)
    doc = Document()
    doc.add_heading(title, 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            cells[i].text = str(row[c])

    doc.save(out_path)


def make_txt_from_csv(csv_path: Path, out_path: Path, title: str) -> None:
    df = pd.read_csv(csv_path)
    lines = [title, "=" * len(title), ""]
    for _, row in df.iterrows():
        parts = [f"{col}={row[col]}" for col in df.columns]
        lines.append(" | ".join(parts))
    out_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", default="data/raw")
    parser.add_argument("--out", default="data/samples")
    args = parser.parse_args()

    raw = Path(args.data)
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)

    pairs = [
        ("settlements.csv", "Razorpay Settlements"),
        ("bank.csv", "Bank Statement"),
    ]
    for fname, title in pairs:
        csv = raw / fname
        if not csv.exists():
            print(f"Skip {csv} (missing)")
            continue
        stem = fname.replace(".csv", "")
        import shutil

        shutil.copy(csv, out / fname)
        make_excel(csv, out / f"{stem}.xlsx")
        make_pdf_from_csv(csv, out / f"{stem}.pdf", title)
        make_docx_from_csv(csv, out / f"{stem}.docx", title)
        make_txt_from_csv(csv, out / f"{stem}.txt", title)
        print(f"Created samples for {stem} -> {out}")


if __name__ == "__main__":
    main()
