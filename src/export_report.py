"""Export reconciliation results to PDF and Word."""

from __future__ import annotations

import json
import zipfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from fpdf import FPDF
from fpdf.enums import XPos, YPos


def _safe(text: object) -> str:
    s = str(text) if text is not None else ""
    return s.encode("latin-1", errors="replace").decode("latin-1")


def build_pdf_report(
    stats: dict,
    matches: list[dict],
    exceptions: list[dict],
    meta: dict | None = None,
) -> bytes:
    meta = meta or {}
    pdf = FPDF(format="A4")
    pdf.set_margins(14, 14, 14)
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    width = pdf.epw

    pdf.set_fill_color(12, 92, 79)
    pdf.rect(0, 0, 210, 32, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_xy(14, 10)
    pdf.cell(width, 10, _safe("SettleMatch Reconciliation Report"))
    pdf.set_font("Helvetica", "", 9)
    pdf.set_xy(14, 22)
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    pdf.cell(width, 6, _safe(f"Generated {ts}  |  Razorpay AI Buildathon Track 04"))

    pdf.set_text_color(30, 30, 30)
    pdf.set_xy(pdf.l_margin, 40)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(width, 8, _safe("Executive summary"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 10)
    lines = [
        f"Match rate: {stats.get('match_rate', 0) * 100:.1f}%",
        f"Recovery rate: {stats.get('recovery_rate', 0) * 100:.1f}%",
        f"Matched: {stats.get('matched', 0)} / {stats.get('total_settlements', 0)}",
        f"Money matched: INR {stats.get('money_matched_inr', 0):,.2f}",
        f"Money at risk: INR {stats.get('money_at_risk_inr', 0):,.2f}",
        f"Collisions detected: {stats.get('collisions_detected', 0)}",
        f"Rule matches: {stats.get('rule_matches', 0)}  |  AI matches: {stats.get('ai_matches', 0)}",
        f"Exceptions (refused): {stats.get('exceptions', 0)}",
        f"AI mode: {stats.get('ai_mode', 'n/a')}  |  Gate: {stats.get('confidence_threshold', 0.85)}",
    ]
    if meta.get("settlements_source"):
        lines.append(f"Settlements source: {meta['settlements_source']}")
    if meta.get("bank_source"):
        lines.append(f"Bank source: {meta['bank_source']}")

    for line in lines:
        pdf.multi_cell(width, 6, _safe(line))

    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 12)
    pdf.multi_cell(width, 8, _safe("Top matches (sample)"))
    pdf.set_font("Helvetica", "", 9)
    for m in matches[:12]:
        pdf.multi_cell(
            width,
            5,
            _safe(
                f"{m.get('settlement_id')} -> {m.get('bank_txn_id')} "
                f"({m.get('method')}, conf={m.get('confidence')})"
            ),
        )

    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 12)
    pdf.multi_cell(width, 8, _safe("Exceptions (unresolved)"))
    pdf.set_font("Helvetica", "", 9)
    if not exceptions:
        pdf.multi_cell(width, 5, _safe("None - all settlements matched."))
    else:
        for ex in exceptions[:15]:
            pdf.multi_cell(
                width,
                5,
                _safe(
                    f"{ex.get('settlement_id')} | INR {ex.get('amount')} | {ex.get('reason')}"
                ),
            )

    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(
        width,
        5,
        _safe("Rules first. AI second. Gate always. Every decision is auditable."),
    )

    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1")


def build_word_report(
    stats: dict,
    matches: list[dict],
    exceptions: list[dict],
    meta: dict | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = doc.add_heading("SettleMatch Reconciliation Report", 0)
    title.runs[0].font.size = Pt(22)

    doc.add_paragraph(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    doc.add_heading("Executive summary", level=1)
    for key in [
        ("Match rate", f"{stats.get('match_rate', 0) * 100:.1f}%"),
        ("Matched", f"{stats.get('matched')} / {stats.get('total_settlements')}"),
        ("Money matched (INR)", f"{stats.get('money_matched_inr', 0):,.2f}"),
        ("Rule / AI", f"{stats.get('rule_matches')} / {stats.get('ai_matches')}"),
        ("Exceptions", str(stats.get("exceptions"))),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{key[0]}: ").bold = True
        p.add_run(key[1])

    if meta:
        if meta.get("settlements_source"):
            doc.add_paragraph(f"Settlements: {meta['settlements_source']}")
        if meta.get("bank_source"):
            doc.add_paragraph(f"Bank: {meta['bank_source']}")

    doc.add_heading("Matches (sample)", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Settlement", "Bank", "Method", "Confidence"]):
        hdr[i].text = h
    for m in matches[:20]:
        row = table.add_row().cells
        row[0].text = str(m.get("settlement_id", ""))
        row[1].text = str(m.get("bank_txn_id", ""))
        row[2].text = str(m.get("method", ""))
        row[3].text = str(m.get("confidence", ""))

    doc.add_heading("Exceptions", level=1)
    if not exceptions:
        doc.add_paragraph("None.")
    else:
        for ex in exceptions:
            doc.add_paragraph(
                f"{ex.get('settlement_id')} | INR {ex.get('amount')} | {ex.get('reason')}",
                style="List Bullet",
            )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def write_pdf_to_disk(stats, matches, exceptions, path: str | Path, meta=None) -> Path:
    path = Path(path)
    path.write_bytes(build_pdf_report(stats, matches, exceptions, meta))
    return path


def build_audit_zip(
    stats: dict,
    matches: list[dict],
    exceptions: list[dict],
    collisions: list[dict],
    audit_events: list,
    meta: dict | None = None,
) -> bytes:
    """One-click audit pack — PDF, Word, JSON, CSV inside ZIP."""
    import pandas as pd

    meta = meta or {}
    buf = BytesIO()
    pdf_bytes = build_pdf_report(stats, matches, exceptions, meta)
    docx_bytes = build_word_report(stats, matches, exceptions, meta)

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("SettleMatch_report.pdf", pdf_bytes)
        zf.writestr("SettleMatch_report.docx", docx_bytes)
        zf.writestr("report.json", json.dumps(stats, indent=2))
        zf.writestr("matches.csv", pd.DataFrame(matches).to_csv(index=False))
        zf.writestr("exceptions.csv", pd.DataFrame(exceptions).to_csv(index=False))
        zf.writestr("collisions.json", json.dumps(collisions, indent=2))
        audit_lines = [
            json.dumps(e.__dict__ if hasattr(e, "__dict__") else e)
            for e in audit_events
        ]
        zf.writestr("audit_log.jsonl", "\n".join(audit_lines))
        zf.writestr(
            "README.txt",
            "SettleMatch Audit Pack\n"
            f"Generated {datetime.now(timezone.utc).isoformat()}\n"
            "Files: PDF report, Word report, matches, exceptions, collisions, audit log\n",
        )
    return buf.getvalue()
