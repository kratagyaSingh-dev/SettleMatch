"""Multi-format document extraction: CSV, Excel, PDF, Word, TXT."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

import pandas as pd

SETTLEMENT_MAP = {
    "settlement_id": ["settlement_id", "settlement id", "settle_id", "settlement", "id"],
    "payment_id": ["payment_id", "payment id", "pay_id", "payment"],
    "amount": ["amount", "amt", "value", "settlement_amount", "net_amount"],
    "currency": ["currency", "curr", "ccy"],
    "utr": ["utr", "utr_no", "utr number", "reference", "ref_no", "bank_ref"],
    "settled_at": ["settled_at", "settled at", "settlement_date", "date", "settled_date"],
    "status": ["status", "state", "settlement_status"],
}

BANK_MAP = {
    "bank_txn_id": ["bank_txn_id", "txn_id", "transaction_id", "id", "ref"],
    "amount": ["amount", "amt", "credit", "value", "transaction_amount"],
    "narration": ["narration", "description", "particulars", "remarks", "details"],
    "value_date": ["value_date", "value date", "date", "txn_date", "transaction_date"],
    "utr": ["utr", "utr_no", "reference", "ref_no", "cheque_no"],
}

SETTLEMENT_COLS = list(SETTLEMENT_MAP.keys())
BANK_COLS = list(BANK_MAP.keys())

SUPPORTED = {".csv", ".xlsx", ".xls", ".pdf", ".docx", ".txt"}


def required_columns(kind: str) -> list[str]:
    return SETTLEMENT_COLS if kind == "settlements" else BANK_COLS


@dataclass
class ExtractResult:
    dataframe: pd.DataFrame
    source_format: str
    rows_found: int
    columns_mapped: dict[str, str]
    warnings: list[str]


def _norm_key(name: str) -> str:
    return re.sub(r"[^a-z0-9]", "_", str(name).lower().strip())


def _map_columns(df: pd.DataFrame, schema: dict[str, list[str]]) -> tuple[pd.DataFrame, dict[str, str], list[str]]:
    warnings: list[str] = []
    col_lookup = {_norm_key(c): c for c in df.columns}
    mapped: dict[str, str] = {}
    out = pd.DataFrame()

    for target, aliases in schema.items():
        found = None
        for alias in aliases:
            key = _norm_key(alias)
            if key in col_lookup:
                found = col_lookup[key]
                break
        if found is None:
            for alias in aliases:
                key = _norm_key(alias)
                for nk, orig in col_lookup.items():
                    if key in nk or nk in key:
                        found = orig
                        break
                if found:
                    break
        if found:
            out[target] = df[found]
            mapped[target] = found
        elif target in ("currency", "status", "payment_id", "narration", "utr"):
            if target == "currency":
                out[target] = "INR"
            elif target == "status":
                out[target] = "settled"
            elif target == "payment_id":
                out[target] = out.get("settlement_id", pd.Series(dtype=str)).apply(
                    lambda x: f"pay_{x}" if pd.notna(x) else ""
                )
            elif target == "narration":
                out[target] = ""
            elif target == "utr":
                out[target] = ""
            warnings.append(f"Column '{target}' not found — filled default")
        else:
            warnings.append(f"Required column '{target}' could not be mapped")

    return out, mapped, warnings


def _read_csv(raw: bytes) -> pd.DataFrame:
    return pd.read_csv(io.BytesIO(raw))


def _read_excel(raw: bytes) -> pd.DataFrame:
    xl = pd.ExcelFile(io.BytesIO(raw))
    best = None
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if df.empty:
            continue
        if best is None or len(df) > len(best):
            best = df
    if best is None:
        raise ValueError("No data found in Excel workbook")
    return best


def _read_pdf(raw: bytes) -> pd.DataFrame:
    import pdfplumber

    text_lines: list[str] = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            text_lines.extend(text.splitlines())

    structured = [ln for ln in text_lines if "=" in ln and "|" in ln]
    if len(structured) >= 3:
        return _read_structured_txt(text_lines)

    # Fallback: native PDF tables (legacy samples)
    frames: list[pd.DataFrame] = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                if not table or len(table) < 2:
                    continue
                header = [str(c or "").strip() for c in table[0]]
                if not any(header):
                    continue
                rows = []
                for row in table[1:]:
                    cells = [str(c or "").strip() for c in row]
                    if any(cells):
                        rows.append(cells)
                if not rows:
                    continue
                width = len(header)
                clean_rows = [r[:width] + [""] * max(0, width - len(r)) for r in rows]
                frames.append(pd.DataFrame(clean_rows, columns=header))

    if frames:
        return _merge_tables(frames)

    return _parse_text_lines(text_lines)


def _merge_tables(frames: list[pd.DataFrame]) -> pd.DataFrame:
    """Merge PDF page tables that share the same header row."""
    if len(frames) == 1:
        return frames[0]

    norm_headers = [[_norm_key(c) for c in df.columns] for df in frames]
    merged: pd.DataFrame | None = None
    ref_header: list[str] | None = None

    for df, hdr in zip(frames, norm_headers):
        if merged is None:
            merged = df.copy()
            ref_header = hdr
            continue
        if hdr == ref_header:
            merged = pd.concat([merged, df], ignore_index=True)
        elif len(df) > len(merged):
            merged = df.copy()
            ref_header = hdr

    return merged if merged is not None else frames[0]


def _read_docx(raw: bytes) -> pd.DataFrame:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    best = None
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        df = pd.DataFrame(rows[1:], columns=rows[0])
        if best is None or len(df) > len(best):
            best = df
    if best is not None:
        return best
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    return _parse_text_lines(lines)


def _read_txt(raw: bytes) -> pd.DataFrame:
    text = raw.decode("utf-8", errors="replace")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    structured = [ln for ln in lines if "=" in ln and "|" in ln]
    if len(structured) >= 3:
        return _read_structured_txt(lines)
    if lines and "," in lines[0]:
        try:
            return pd.read_csv(io.BytesIO(raw))
        except Exception:
            pass
    return _parse_text_lines(lines)


def _read_structured_txt(lines: list[str]) -> pd.DataFrame:
    rows: list[dict[str, str]] = []
    for line in lines:
        if "=" not in line or "|" not in line:
            continue
        if line.startswith("="):
            continue
        row: dict[str, str] = {}
        for part in line.split("|"):
            part = part.strip()
            if "=" not in part:
                continue
            k, _, v = part.partition("=")
            val = v.strip()
            if val.lower() in ("nan", "none", ""):
                val = ""
            row[k.strip()] = val
        if row:
            rows.append(row)
    if not rows:
        raise ValueError("Could not parse structured TXT/PDF text.")
    return pd.DataFrame(rows)


def _parse_text_lines(lines: list[str]) -> pd.DataFrame:
    """Regex fallback when PDF/Word has no tables."""
    rows = []
    utr_re = re.compile(r"UTR[A-Z0-9]{6,}", re.I)
    amt_re = re.compile(r"(?:INR|Rs\.?|₹)?\s*([\d,]+\.?\d*)")
    date_re = re.compile(r"(\d{4}-\d{2}-\d{2}|\d{2}[/-]\d{2}[/-]\d{4})")

    for i, line in enumerate(lines):
        line = line.strip()
        if len(line) < 8:
            continue
        utr_m = utr_re.search(line)
        amt_m = amt_re.search(line)
        date_m = date_re.search(line)
        if not amt_m:
            continue
        amount = float(amt_m.group(1).replace(",", ""))
        utr = utr_m.group(0).upper() if utr_m else ""
        dt = date_m.group(1) if date_m else ""
        rows.append(
            {
                "amount": amount,
                "utr": utr,
                "date": dt,
                "narration": line[:120],
                "id": f"ROW-{i+1}",
            }
        )

    if not rows:
        raise ValueError("Could not extract rows from text. Try CSV or Excel.")
    return pd.DataFrame(rows)


def extract_bytes(raw: bytes, filename: str, kind: str) -> ExtractResult:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported format '{ext}'. Use: {', '.join(sorted(SUPPORTED))}")

    if ext == ".csv":
        df = _read_csv(raw)
    elif ext in (".xlsx", ".xls"):
        df = _read_excel(raw)
    elif ext == ".pdf":
        df = _read_pdf(raw)
    elif ext == ".docx":
        df = _read_docx(raw)
    elif ext == ".txt":
        df = _read_txt(raw)
    else:
        raise ValueError(f"Unsupported: {ext}")

    schema = SETTLEMENT_MAP if kind == "settlements" else BANK_MAP
    mapped_df, col_map, warnings = _map_columns(df, schema)

    # Text-parse fallback column fixes (legacy paths)
    if kind == "settlements":
        if "settlement_id" not in mapped_df.columns and "id" in df.columns:
            mapped_df["settlement_id"] = df[_find_raw(df, "id")]
        if "settled_at" not in mapped_df.columns and "date" in df.columns:
            mapped_df["settled_at"] = df[_find_raw(df, "date")]
    mapped_df = _finalize(mapped_df, kind, df)
    mapped_df = _strip_to_schema(mapped_df, kind)
    if mapped_df.empty:
        raise ValueError("Extracted table is empty after mapping.")

    extra_in_output = set(mapped_df.columns) - set(required_columns(kind))
    if extra_in_output:
        raise ValueError(f"Internal error: unexpected columns {extra_in_output}")

    return ExtractResult(
        dataframe=mapped_df,
        source_format=ext.lstrip("."),
        rows_found=len(mapped_df),
        columns_mapped=col_map,
        warnings=warnings,
    )


def _finalize(mapped: pd.DataFrame, kind: str, raw: pd.DataFrame) -> pd.DataFrame:
    """Ensure all required ingest columns exist."""
    out = mapped.copy()
    if out.empty and not raw.empty:
        out = pd.DataFrame(index=raw.index)
    raw_cols = {_norm_key(c): c for c in raw.columns}

    if kind == "settlements":
        if "settlement_id" not in out.columns:
            if "id" in raw_cols:
                out["settlement_id"] = raw[raw_cols["id"]]
            else:
                out["settlement_id"] = [f"SET-{i+1:03d}" for i in range(len(out))]
        if "settled_at" not in out.columns and "date" in raw_cols:
            out["settled_at"] = raw[raw_cols["date"]]
        if "payment_id" not in out.columns:
            out["payment_id"] = out["settlement_id"].astype(str).apply(
                lambda x: f"pay_{x}" if x else ""
            )
        if "currency" not in out.columns:
            out["currency"] = "INR"
        if "status" not in out.columns:
            out["status"] = "settled"
        if "utr" not in out.columns:
            out["utr"] = ""
    else:
        if "bank_txn_id" not in out.columns:
            if "id" in raw_cols:
                out["bank_txn_id"] = raw[raw_cols["id"]]
            else:
                out["bank_txn_id"] = [f"BK-{i+1:04d}" for i in range(len(out))]
        if "value_date" not in out.columns and "date" in raw_cols:
            out["value_date"] = raw[raw_cols["date"]]
        if "narration" not in out.columns:
            out["narration"] = ""
        if "utr" not in out.columns:
            out["utr"] = ""

    schema = SETTLEMENT_MAP if kind == "settlements" else BANK_MAP
    for col in schema:
        if col not in out.columns:
            out[col] = ""

    return _strip_to_schema(out, kind)


def _strip_to_schema(df: pd.DataFrame, kind: str) -> pd.DataFrame:
    """Keep ONLY required columns — drop any extra from PDF/Excel/Word."""
    cols = required_columns(kind)
    out = df.copy()
    for c in cols:
        if c not in out.columns:
            out[c] = ""
    return out[cols].copy()


def _find_raw(df: pd.DataFrame, alias: str) -> str:
    for c in df.columns:
        if _norm_key(c) == _norm_key(alias):
            return c
    return df.columns[0]


def save_extracted(result: ExtractResult, path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    result.dataframe.to_csv(path, index=False)
    return path
